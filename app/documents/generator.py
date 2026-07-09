"""Document generation — creates PDF and DOCX submission files with professional formatting."""

import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch, mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from config.settings import GENERATED_DIR
from app.documents.templates import get_student_info

logger = logging.getLogger(__name__)

# ─── Colors ─────────────────────────────────────────────────────────────

C_PRIMARY = colors.HexColor("#1a237e")
C_ACCENT = colors.HexColor("#283593")
C_CODE_BG = colors.HexColor("#f5f5f5")
C_CODE_BORDER = colors.HexColor("#e0e0e0")
C_IO_BG = colors.HexColor("#e8f5e9")
C_IO_BORDER = colors.HexColor("#a5d6a7")
C_INPUT_BG = colors.HexColor("#e3f2fd")
C_INPUT_BORDER = colors.HexColor("#90caf9")
C_OUTPUT_BG = colors.HexColor("#fff3e0")
C_OUTPUT_BORDER = colors.HexColor("#ffcc80")
C_SEPARATOR = colors.HexColor("#bbdefb")
C_HEADER_BG = colors.HexColor("#e8eaf6")
C_HEADER_BORDER = colors.HexColor("#c5cae9")

# ─── Styles ──────────────────────────────────────────────────────────────

def _build_styles() -> dict:
    """Return a dict of reusable ParagraphStyle objects."""
    base = getSampleStyleSheet()
    return {
        "header_label": ParagraphStyle(
            "HdrLabel", parent=base["Normal"],
            fontSize=10, leading=14, textColor=C_PRIMARY,
            spaceAfter=0, spaceBefore=0,
        ),
        "header_value": ParagraphStyle(
            "HdrValue", parent=base["Normal"],
            fontSize=10, leading=14,
            spaceAfter=0, spaceBefore=0,
        ),
        "title": ParagraphStyle(
            "DocTitle", parent=base["Title"],
            fontSize=18, leading=24, alignment=TA_CENTER,
            textColor=C_PRIMARY, spaceAfter=16, spaceBefore=8,
        ),
        "heading1": ParagraphStyle(
            "DocH1", parent=base["Heading1"],
            fontSize=16, leading=22, textColor=C_PRIMARY,
            spaceBefore=18, spaceAfter=8,
        ),
        "heading2": ParagraphStyle(
            "DocH2", parent=base["Heading2"],
            fontSize=14, leading=20, textColor=C_ACCENT,
            spaceBefore=14, spaceAfter=6,
        ),
        "heading3": ParagraphStyle(
            "DocH3", parent=base["Heading3"],
            fontSize=12, leading=18, textColor=C_PRIMARY,
            spaceBefore=10, spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "DocBody", parent=base["Normal"],
            fontSize=11, leading=16, alignment=TA_JUSTIFY,
            spaceAfter=6, spaceBefore=0,
        ),
        "bullet": ParagraphStyle(
            "DocBullet", parent=base["Normal"],
            fontSize=11, leading=16, leftIndent=20, bulletIndent=8,
            spaceAfter=3, spaceBefore=0,
        ),
        "code": ParagraphStyle(
            "DocCode", parent=base["Code"],
            fontSize=8.5, leading=12, backColor=C_CODE_BG,
            leftIndent=8, rightIndent=8, spaceAfter=0, spaceBefore=0,
            fontName="Courier",
        ),
        "io_label": ParagraphStyle(
            "IOLabel", parent=base["Normal"],
            fontSize=10, leading=14, textColor=colors.HexColor("#2e7d32"),
            spaceAfter=2, spaceBefore=0, fontName="Helvetica-Bold",
        ),
        "io_content": ParagraphStyle(
            "IOContent", parent=base["Normal"],
            fontSize=9.5, leading=13, fontName="Courier",
            spaceAfter=0, spaceBefore=0,
        ),
    }


# ─── Path helper ────────────────────────────────────────────────────────

def _build_absolute_path(assignment_id: str, title: str, ext: str) -> Path:
    safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)
    safe_title = safe_title.strip().replace(" ", "_")[:60]
    return GENERATED_DIR / f"{assignment_id}_{safe_title}.{ext}"


# ─── Header block ───────────────────────────────────────────────────────

def _build_header_table(info: dict, subject: str, today: str, styles: dict) -> Table:
    """Create a bordered table with student info header."""
    rows = []
    for label, value in [
        ("Student Name", info["name"]),
        ("Register No", info["reg_no"]),
        ("Subject", subject),
        ("Date", today),
    ]:
        label_cell = Paragraph(f"<b>{label}</b>", styles["header_label"])
        value_cell = Paragraph(value, styles["header_value"])
        rows.append([label_cell, value_cell])

    col_widths = [95, 380]
    t = Table(rows, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1, C_HEADER_BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, C_HEADER_BORDER),
        ("BACKGROUND", (0, 0), (0, -1), C_HEADER_BG),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ]))
    return t


# ─── Code block renderer ────────────────────────────────────────────────

def _render_code_block(lines: list[str], styles: dict) -> list:
    """Render a code block as a bordered table with grey background."""
    elements = []
    code_lines = []
    in_code = False
    for line in lines:
        if line.startswith("```"):
            if in_code:
                break
            in_code = True
            continue
        if in_code:
            code_lines.append(line)

    if not code_lines:
        return elements

    cell_lines = [Paragraph(ln or " ", styles["code"]) for ln in code_lines]
    cell = Table([[cl] for cl in cell_lines], colWidths=[460])
    cell.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1, C_CODE_BORDER),
        ("BACKGROUND", (0, 0), (-1, -1), C_CODE_BG),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))

    elements.append(Spacer(1, 4))
    elements.append(cell)
    elements.append(Spacer(1, 4))
    return elements


# ─── I/O section renderer ───────────────────────────────────────────────

def _render_io_section(label: str, content: str, bg_color, border_color, label_style_label: str, styles: dict) -> list:
    """Render an Input/Output section with colored background and border."""
    io_style = styles.get(label_style_label, styles["io_content"])
    lines = content.strip().split("\n")
    cell_lines = [Paragraph(ln or " ", io_style) for ln in lines]
    cell = Table([[cl] for cl in cell_lines], colWidths=[460])
    cell.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1.5, border_color),
        ("BACKGROUND", (0, 0), (-1, -1), bg_color),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))

    label_p = Paragraph(f"<b>{label}</b>", styles["io_label"])
    return [label_p, cell, Spacer(1, 6)]


def _detect_and_render_io(lines: list[str], idx: int, styles: dict) -> tuple[list, int]:
    """Check if lines starting at idx contain I/O content. Returns (elements, consumed)."""
    if idx >= len(lines):
        return [], 0

    line = lines[idx].strip()
    # Strip markdown heading markers for matching
    clean_line = re.sub(r"^#+\s*", "", line).rstrip(":").strip()

    io_patterns = {
        "Input": (C_INPUT_BG, C_INPUT_BORDER, "io_content"),
        "Output": (C_OUTPUT_BG, C_OUTPUT_BORDER, "io_content"),
        "Sample Input": (C_INPUT_BG, C_INPUT_BORDER, "io_content"),
        "Sample Output": (C_OUTPUT_BG, C_OUTPUT_BORDER, "io_content"),
        "Example Input": (C_INPUT_BG, C_INPUT_BORDER, "io_content"),
        "Example Output": (C_OUTPUT_BG, C_OUTPUT_BORDER, "io_content"),
    }

    matched_label = None
    for key in io_patterns:
        if clean_line.lower() == key.lower():
            matched_label = key
            break
        if clean_line.lower().startswith(key.lower()):
            matched_label = key
            break

    if not matched_label:
        return [], 0

    bg, border, content_style_key = io_patterns[matched_label]
    content_lines = []
    consumed = 1

    for j in range(idx + 1, len(lines)):
        next_line = lines[j]
        if next_line.startswith("```"):
            code_elements = _render_code_block(lines[j:], styles)
            if code_elements:
                for el in code_elements[1:-1]:
                    content_lines.append(el)
                consumed = j - idx + (len(code_elements) if code_elements else 0)
                break
        elif next_line.strip().startswith("##") or next_line.strip().startswith("#"):
            break
        elif not next_line.strip():
            content_lines.append(Paragraph(" ", styles["body"]))
            consumed += 1
        else:
            content_lines.append(Paragraph(next_line, styles["io_content"]))
            consumed += 1
    else:
        consumed = len(lines) - idx

    cell = Table([[cl] for cl in content_lines], colWidths=[460])
    cell.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 1.5, border),
        ("BACKGROUND", (0, 0), (-1, -1), bg),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))

    label_p = Paragraph(f"<b>{matched_label}</b>", styles["io_label"])
    return [label_p, cell, Spacer(1, 6)], consumed


# ─── Page template (header/footer) ──────────────────────────────────────

def _page_template(canvas, doc):
    """Draw header line and page number on each page."""
    canvas.saveState()
    # Header line
    canvas.setStrokeColor(C_HEADER_BORDER)
    canvas.setLineWidth(0.5)
    canvas.line(inch, A4[1] - 0.5 * inch, A4[0] - inch, A4[1] - 0.5 * inch)
    # Page number
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.grey)
    canvas.drawCentredString(A4[0] / 2, 0.5 * inch, f"Page {doc.page}")
    canvas.restoreState()


# ─── PDF Generation ─────────────────────────────────────────────────────

def generate_pdf(
    assignment_id: str,
    title: str,
    subject: str,
    solution_markdown: str,
) -> Path:
    """Generate a professionally formatted PDF from markdown solution content."""
    out_path = _build_absolute_path(assignment_id, title, "pdf")
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        topMargin=0.7 * inch,
        bottomMargin=0.8 * inch,
        leftMargin=1.0 * inch,
        rightMargin=1.0 * inch,
    )

    styles = _build_styles()
    info = get_student_info()
    today = datetime.now().strftime("%B %d, %Y")
    elements = []

    # ── Header block ────────────────────────────────────────────────
    elements.append(_build_header_table(info, subject, today, styles))
    elements.append(Spacer(1, 14))

    # ── Title ───────────────────────────────────────────────────────
    elements.append(Paragraph(title, styles["title"]))
    elements.append(HRFlowable(
        width="100%", thickness=1.5, color=C_SEPARATOR, spaceAfter=12, spaceBefore=0,
    ))

    # ── Markdown → PDF content ──────────────────────────────────────
    lines = solution_markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip template YAML front matter
        if i == 0 and line == "---":
            while i < len(lines) and not lines[i].strip().startswith("---"):
                i += 1
            i += 1
            continue
        if line.startswith("---") and "date" not in line:
            i += 1
            continue

        # Check for I/O sections first
        io_elements, consumed = _detect_and_render_io(lines, i, styles)
        if consumed > 0:
            elements.extend(io_elements)
            i += consumed
            continue

        # Code blocks
        if line.startswith("```"):
            code_elements = _render_code_block(lines[i:], styles)
            elements.extend(code_elements)
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # Empty line
        if not line:
            elements.append(Spacer(1, 4))
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            elements.append(Paragraph(line[2:], styles["heading1"]))
            i += 1
            continue
        if line.startswith("## "):
            elements.append(Paragraph(line[3:], styles["heading2"]))
            i += 1
            continue
        if line.startswith("### "):
            elements.append(Paragraph(line[4:], styles["heading3"]))
            i += 1
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            elements.append(Paragraph(f"&bull; {text}", styles["bullet"]))
            i += 1
            continue

        # Numbered list
        match = re.match(r"^\d+[.)]\s+", line)
        if match:
            text = line[match.end():]
            num = line[:match.end() - 1]
            elements.append(Paragraph(f"<b>{num}.</b>  {text}", styles["body"]))
            i += 1
            continue

        # Bold/italic markers at line level
        if re.match(r"^\*\*.*\*\*", line):
            elements.append(Paragraph(line, styles["body"]))
            i += 1
            continue

        # Regular paragraph
        elements.append(Paragraph(line, styles["body"]))
        i += 1

    # ── Footer separator ────────────────────────────────────────────
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(
        width="100%", thickness=0.5, color=C_CODE_BORDER, spaceAfter=6, spaceBefore=0,
    ))
    elements.append(Paragraph(
        "Generated by AI-Powered LMS Assistant",
        ParagraphStyle("Footer", fontSize=8, textColor=colors.grey, alignment=TA_CENTER),
    ))

    doc.build(elements, onFirstPage=_page_template, onLaterPages=_page_template)
    logger.info("Generated PDF: %s", out_path)
    return out_path


# ─── DOCX Generation ────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _add_formatted_paragraph(doc, text: str, style_name: str = "Normal", bold: bool = False, size: int = 11, color: str = None):
    """Add a paragraph with formatting. Handles inline bold (**text**)."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (1, 3, 5)])
    return p


def generate_docx(
    assignment_id: str,
    title: str,
    subject: str,
    solution_markdown: str,
) -> Path:
    """Generate a professionally formatted DOCX file from markdown solution content."""
    out_path = _build_absolute_path(assignment_id, title, "docx")
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    info = get_student_info()
    today = datetime.now().strftime("%B %d, %Y")

    # ── Header table ────────────────────────────────────────────────
    header_table = doc.add_table(rows=4, cols=2)
    header_table.style = "Table Grid"
    for idx, (label, value) in enumerate([
        ("Student Name", info["name"]),
        ("Register No", info["reg_no"]),
        ("Subject", subject),
        ("Date", today),
    ]):
        cell_label = header_table.rows[idx].cells[0]
        cell_value = header_table.rows[idx].cells[1]
        cell_label.text = ""
        cell_value.text = ""
        run_l = cell_label.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(10)
        _set_cell_shading(cell_label, "e8eaf6")
        run_v = cell_value.paragraphs[0].add_run(value)
        run_v.font.size = Pt(10)

    doc.add_paragraph()

    # ── Title ───────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(18)
    run_t.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

    # Separator
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sep = p_sep.add_run("─" * 70)
    run_sep.font.color.rgb = RGBColor(0xbb, 0xde, 0xfb)
    run_sep.font.size = Pt(8)

    # ── Content ─────────────────────────────────────────────────────
    in_code_block = False
    code_lines = []
    in_io_block = False
    io_type = ""
    block_io_lines = []

    def flush_code_block():
        nonlocal code_lines, in_code_block
        if code_lines:
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            _set_cell_shading(cell, "f5f5f5")
            cell.text = ""
            for cl in code_lines:
                p = cell.paragraphs[0] if cell.paragraphs[0] == cell.paragraphs[0] else cell.add_paragraph()
                if cl == "":
                    p.add_run(" ")
                else:
                    run = p.add_run(cl)
                    run.font.name = "Courier New"
                    run.font.size = Pt(8.5)
            code_lines = []
            in_code_block = False

    def flush_io_block():
        nonlocal block_io_lines, io_type, in_io_block
        if block_io_lines:
            bg = "e3f2fd" if "input" in io_type.lower() else "fff3e0"
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            _set_cell_shading(cell, bg)
            cell.text = ""
            # Label
            label_run = cell.paragraphs[0].add_run(f"{io_type}")
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0x2e, 0x7d, 0x32)
            cell.paragraphs[0].add_run("\n")
            for bl in block_io_lines:
                if bl == "":
                    cell.paragraphs[0].add_run(" ")
                else:
                    run = cell.paragraphs[0].add_run(bl + "\n")
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
            block_io_lines = []
            io_type = ""
            in_io_block = False

    lines = solution_markdown.split("\n")
    for line in lines:
        stripped = line.strip()

        # Skip YAML front matter
        if stripped == "---":
            continue

        # Code block handling
        if stripped.startswith("```"):
            if in_code_block:
                flush_code_block()
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # I/O detection
        io_keywords = ["input", "output", "sample input", "sample output",
                       "example input", "example output"]
        clean_stripped = re.sub(r"^#+\s*", "", stripped).rstrip(":").strip().lower()
        if clean_stripped in io_keywords or \
           any(clean_stripped.startswith(k) for k in io_keywords):
            if in_io_block:
                flush_io_block()
            io_type = stripped.rstrip(":")
            in_io_block = True
            continue

        if in_io_block:
            block_io_lines.append(line)
            continue

        # Empty line
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x28, 0x35, 0x93)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
            continue

        # Bullets
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        # Numbered
        match = re.match(r"^(\d+)[.)]\s+(.*)", stripped)
        if match:
            p = doc.add_paragraph()
            run_num = p.add_run(f"{match.group(1)}. ")
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_num.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
            run_rest = p.add_run(match.group(2))
            run_rest.font.size = Pt(11)
            continue

        # Regular paragraph with inline bold
        _add_formatted_paragraph(doc, stripped, size=11)

    # Flush any remaining blocks
    flush_code_block()
    flush_io_block()

    doc.save(str(out_path))
    logger.info("Generated DOCX: %s", out_path)
    return out_path
